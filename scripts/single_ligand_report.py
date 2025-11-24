import os
from docx import Document
from datetime import datetime

# ✅ Input files
POSE_FILE = "docking/poses/pose0_out.pdbqt"
IMAGE_FILE = "results/images/pose0.png"

# ✅ Output file
REPORT_FILE = "results/report/6LU7_single_ligand_report.docx"

# ✅ Make sure output folder exists
os.makedirs("results/report", exist_ok=True)

# ✅ Extract binding affinity
def extract_affinity(pdbqt_file):
    with open(pdbqt_file, "r") as f:
        for line in f:
            if line.strip().startswith("REMARK VINA RESULT"):
                parts = line.split()
                return float(parts[3])
    return None

affinity = extract_affinity(POSE_FILE)

# ✅ Build Word document
doc = Document()
doc.add_heading("SARS-CoV-2 Mpro (6LU7) Docking Report", level=1)

doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
doc.add_paragraph("Target Protein: SARS-CoV-2 Main Protease (Mpro, PDB ID: 6LU7)")
doc.add_paragraph("Ligand: pose0")

doc.add_heading("1. Binding Affinity", level=2)
doc.add_paragraph(f"Binding Affinity (kcal/mol): {affinity}")

doc.add_heading("2. Binding Pose Visualization", level=2)
if os.path.exists(IMAGE_FILE):
    doc.add_picture(IMAGE_FILE)
else:
    doc.add_paragraph("⚠ Image not found. Run auto_visualize.py first.")

doc.add_heading("3. Summary", level=2)
doc.add_paragraph(
    "This result provides an initial insight into the ligand’s predicted "
    "interaction with SARS-CoV-2 Mpro. More negative affinity values indicate "
    "stronger predicted binding."
)

doc.save(REPORT_FILE)

print("✅ Report Generated:")
print(REPORT_FILE)
